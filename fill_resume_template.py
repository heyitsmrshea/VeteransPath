from docx import Document

def fill_resume_template(
    template_path,
    summary,
    experience_blocks,
    full_name,
    email,
    phone,
    linkedin,
    clearance,
    education,
    certifications
):
    doc = Document(template_path)

    for p in doc.paragraphs:
        if "{{FULL_NAME}}" in p.text:
            p.text = p.text.replace("{{FULL_NAME}}", full_name)
        if "{{EMAIL}}" in p.text:
            p.text = p.text.replace("{{EMAIL}}", email)
        if "{{PHONE}}" in p.text:
            p.text = p.text.replace("{{PHONE}}", phone)
        if "{{LINKEDIN}}" in p.text:
            p.text = p.text.replace("{{LINKEDIN}}", linkedin)
        if "{{CLEARANCE}}" in p.text:
            p.text = p.text.replace("{{CLEARANCE}}", clearance)
        if "{{SUMMARY}}" in p.text:
            p.text = p.text.replace("{{SUMMARY}}", summary)
        if "{{EDUCATION}}" in p.text:
            p.text = p.text.replace("{{EDUCATION}}", education)
        if "{{CERTIFICATIONS}}" in p.text:
            p.text = p.text.replace("{{CERTIFICATIONS}}", ", ".join(certifications))

    for i, p in enumerate(doc.paragraphs):
        if "{{EXPERIENCE_BLOCK}}" in p.text:
            parent = p._element.getparent()
            idx = parent.index(p._element)
            parent.remove(p._element)

            for job in experience_blocks:
                role = job.get("title", "")
                unit = job.get("unit", "")
                loc = job.get("location", "")
                start = job.get("start_date", "")
                end = job.get("end_date", "")
                bullets = job.get("bullets", [])

                header = f"{role} — {unit} — {loc} ({start} to {end})"
                doc.paragraphs.insert(i, doc.add_paragraph(header, style="Heading 2"))
                i += 1
                for bullet in bullets:
                    doc.paragraphs.insert(i, doc.add_paragraph(f"• {bullet}", style="List Bullet"))
                    i += 1
            break

    return doc
