import openai
import docx2txt
import PyPDF2
import os
from docx import Document
from datetime import datetime

openai.api_key = os.getenv("OPENAI_API_KEY")

def translate_fitrep_to_resume(text, target_role):
    prompt = f"""
    You are a resume-writing assistant. Translate the following military FITREP into three strong resume bullet points for a role as a {target_role}.
    
    Break it into:
    1. High-Level Overview of the mission/role
    2. Key responsibilities or tools used
    3. Measurable impact and outcome
    
    FITREP:
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt + text}],
            temperature=0.5
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {str(e)}"

def generate_professional_summary(name, rank, mos, branch, years, target_role):
    prompt = f"""
    Write a professional summary for a military veteran named {name}, who served {years} years in the {branch}, held the rank of {rank}, and MOS {mos}. They are now targeting a civilian career in {target_role}. The tone should be confident, civilian-friendly, and tailored for a resume.
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error generating summary: {str(e)}"

def parse_uploaded_resume(uploaded_file):
    if uploaded_file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(uploaded_file)
        return "\n".join([page.extract_text() for page in reader.pages])
    elif uploaded_file.name.endswith(".docx"):
        return docx2txt.process(uploaded_file)
    else:
        return "Unsupported file type. Please upload a .pdf or .docx file."

def parse_uploaded_fitrep(uploaded_file):
    return parse_uploaded_resume(uploaded_file)

def generate_bullets_for_jobs(jobs, target_role):
    all_bullets = []
    for job in jobs:
        role = job.get("role", "")
        responsibilities = job.get("responsibilities", "")
        outcomes = job.get("outcomes", "")
        prompt = f"""
        Convert the following information into 3 strong resume bullet points for a {target_role} role:
        
        1. What was your role or mission context? {role}
        2. Responsibilities or tools used? {responsibilities}
        3. Measurable outcomes or mission impact? {outcomes}
        
        Write this as:
        • Overview of mission/role
        • Key responsibilities or tools used
        • Measurable impact and results
        """
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.4
            )
            bullets = response.choices[0].message.content.strip()
            all_bullets.append(bullets)
        except Exception as e:
            all_bullets.append(f"Error generating bullets: {str(e)}")
    return all_bullets

def generate_word_resume(user_data, bullet_points):
    doc = Document()

    # Name & Contact
    doc.add_heading(user_data.get("name", "Full Name"), 0)
    doc.add_paragraph(f"{user_data.get('city_state', '')} | {user_data.get('phone', '')} | {user_data.get('email', '')}")
    if user_data.get("linkedin"):
        doc.add_paragraph(user_data['linkedin'])
    if user_data.get("clearance"):
        doc.add_paragraph(f"Clearance: {user_data['clearance']}")

    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(user_data.get("summary", ""))

    # Experience
    doc.add_heading("Experience", level=1)
    for i, job in enumerate(user_data.get("jobs", [])):
        title = job.get("job_title", "")
        unit = job.get("unit", "")
        location = job.get("location", "")
        start = job.get("start_date", "")
        end = job.get("end_date", "")

        para = doc.add_paragraph()
        para.add_run(f"{title}, {unit}, {location}\n").bold = True
        para.add_run(f"{start} – {end}")
        doc.add_paragraph(bullet_points[i])

    # Education
    doc.add_heading("Education & Certifications", level=1)
    if user_data.get("education"):
        doc.add_paragraph(user_data["education"])
    for cert in user_data.get("certifications", []):
        doc.add_paragraph(cert)

    # Save docx
    filename = f"Generated_Resume_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    doc.save(filename)
    return filename
