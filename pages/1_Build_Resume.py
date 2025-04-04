import streamlit as st
import openai
import docx
import io
from docx import Document
from docx.shared import Pt
import re
from datetime import datetime

# Set up the Streamlit app
st.set_page_config(page_title="Build My Resume", layout="wide")

# Retrieve OpenAI API key from Streamlit secrets
openai.api_key = st.secrets["OPENAI_API_KEY"]

# Function to clean input text
def clean_input(text):
    return text.replace("\n", " ").replace("•", "-").strip()

# Function to clean and format raw job data with GPT
def format_bullets(raw_text):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a resume writer. Convert raw job input into professional resume bullets."},
            {"role": "user", "content": raw_text}
        ]
    )
    return response['choices'][0]['message']['content']

# Function to extract date from the job experience entry
def extract_date(entry):
    try:
        match = re.search(r'(\w+ \d{4})\s*-\s*(\w+ \d{4}|Present)', entry.get('dates', ''))
        if match:
            end = match.group(2).replace("Present", datetime.now().strftime("%b %Y"))
            return datetime.strptime(end, "%b %Y")
    except:
        pass
    return datetime.min  # Return a default "earliest possible" date if parsing fails

# Utility to remove all content after a given heading
def remove_section(doc, heading_text):
    found = False
    to_remove = []
    for p in doc.paragraphs:
        if heading_text.lower() in p.text.lower():
            found = True
            continue
        if found and p.style.name.startswith("Heading"):
            break
        if found:
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)

# Helper to add bold paragraph
def add_bold_paragraph(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)

# Function to replace text while preserving formatting
def replace_placeholder_with_format(para, placeholder, replacement_text):
    # Iterate through each run in the paragraph (a 'run' is a section of text with the same formatting)
    for run in para.runs:
        if placeholder in run.text:
            # Replace placeholder with the new text
            run.text = run.text.replace(placeholder, replacement_text)

# Function to generate the resume from template
def generate_resume_from_template():
    # Load the Word document template
    doc = Document("Project Manager Template.docx")

    # Replace header placeholders while keeping the format intact
    for para in doc.paragraphs:
        if "[Full Name]" in para.text:
            replace_placeholder_with_format(para, "[Full Name]", st.session_state.full_name)
        elif "[Job Title]" in para.text:
            replace_placeholder_with_format(para, "[Job Title]", st.session_state.job_target)
        elif "[City, State]" in para.text:
            replace_placeholder_with_format(para, "[City, State]", st.session_state.city_state)  # Custom placeholder for City and State
        elif "[Phone Number]" in para.text:
            replace_placeholder_with_format(para, "[Phone Number]", st.session_state.phone)
        elif "[LinkedIn URL]" in para.text:
            replace_placeholder_with_format(para, "[LinkedIn URL]", st.session_state.linkedin)
        elif "[Email]" in para.text:
            replace_placeholder_with_format(para, "[Email]", st.session_state.email)

    # Add summary
    remove_section(doc, "Summary")
    add_bold_paragraph("Summary")
    doc.add_paragraph(st.session_state.summary)

    # Add job experience
    remove_section(doc, "Professional Experience")
    add_bold_paragraph("Professional Experience")
    for job in st.session_state.job_data:
        doc.add_paragraph(f"{job['title']}, {job['unit']} ({job['dates']})")
        doc.add_paragraph(f"{job['location']}")
        doc.add_paragraph(f"Mission: {job['overview']}")
        doc.add_paragraph("Responsibilities:")
        for line in job["responsibilities"].splitlines():
            doc.add_paragraph(f"• {line.strip()}")
        doc.add_paragraph(f"Impact: {job['impact']}")

    # Add education
    remove_section(doc, "Education")
    if st.session_state.education:
        add_bold_paragraph("Education")
        doc.add_paragraph(st.session_state.education)

    # Add certifications
    remove_section(doc, "Certifications")
    if st.session_state.certifications:
        add_bold_paragraph("Certifications")
        for cert in st.session_state.certifications.strip().split("\n"):
            doc.add_paragraph(f"• {cert}")

    # Add skills
    remove_section(doc, "Skills")
    if st.session_state.skills:
        add_bold_paragraph("Skills")
        doc.add_paragraph(", ".join(st.session_state.skills.split(", ")))

    # Prepare for download
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI for data entry
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'positions' not in st.session_state:
    st.session_state.positions = 1
if 'job_data' not in st.session_state:
    st.session_state.job_data = [{} for _ in range(10)]
if 'job_target' not in st.session_state:
    st.session_state.job_target = ""
if 'awards' not in st.session_state:
    st.session_state.awards = ""
if 'volunteer' not in st.session_state:
    st.session_state.volunteer = []
if 'skills' not in st.session_state:
    st.session_state.skills = ""
if 'education' not in st.session_state:
    st.session_state.education = ""
if 'certifications' not in st.session_state:
    st.session_state.certifications = []
if 'summary' not in st.session_state:
    st.session_state.summary = ""

st.title("🧳 Build My Resume")
st.progress((st.session_state.step - 1) / 3)

# Step 1 - Background Information (Updated with Education and Certifications)
if st.session_state.step == 1:
    st.subheader("Step 1 of 3: Your Background")
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.full_name = st.text_input("Full Name", value=st.session_state.get("full_name", ""))
        st.session_state.email = st.text_input("Email", value=st.session_state.get("email", ""))
    with col2:
        st.session_state.phone = st.text_input("Phone Number", value=st.session_state.get("phone", ""))
        st.session_state.linkedin = st.text_input("LinkedIn URL", value=st.session_state.get("linkedin", ""))
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.mos = st.text_input("MOS / Job Code", value=st.session_state.get("mos", ""))
        st.session_state.rank = st.text_input("Last Rank Held", value=st.session_state.get("rank", ""))
        st.session_state.yos = st.number_input("Years of Service", min_value=0, max_value=50, step=1, value=st.session_state.get("yos", 0))
    with col2:
        st.session_state.clearance = st.text_input("Clearance (if any)", value=st.session_state.get("clearance", ""))
        st.session_state.education = st.text_input("Education", value=st.session_state.get("education", ""))
        st.session_state.certifications = st.text_area("Certifications (one per line)", value=st.session_state.get("certifications", ""))
    st.session_state.awards = st.text_area("Awards (one per line)", value=st.session_state.get("awards", ""))
    st.session_state.summary = st.text_area("Summary", value=st.session_state.get("summary", ""))
    if st.button("Next →"):
        st.session_state.step = 2

# Step 2 - Experience, Volunteer, Skills (No Changes)
elif st.session_state.step == 2:
    st.subheader("Step 2 of 3: Experience, Volunteer Work & Skills")
    prev_count = len([entry for entry in st.session_state.job_data if entry])
    new_count = st.number_input("How many jobs do you want to include?", min_value=1, max_value=10, value=st.session_state.positions, step=1)
    if new_count < st.session_state.positions:
        st.session_state.job_data = st.session_state.job_data[:new_count]
    elif new_count > st.session_state.positions:
        st.session_state.job_data.extend([{} for _ in range(new_count - len(st.session_state.job_data))])
    st.session_state.positions = new_count

    for i in range(st.session_state.positions):
        with st.expander(f"Position {i+1}", expanded=True):
            entry = st.session_state.job_data[i]
            entry['title'] = st.text_input("Job Title", value=entry.get('title', ''), key=f"title_{i}")
            entry['unit'] = st.text_input("Unit/Command", value=entry.get('unit', ''), key=f"unit_{i}")
            entry['location'] = st.text_input("Location", value=entry.get('location', ''), key=f"loc_{i}")
            entry['dates'] = st.text_input("Dates (e.g., Jan 2020 - Jan 2023)", value=entry.get('dates', ''), key=f"dates_{i}")
            entry['overview'] = st.text_area("Mission Overview", value=entry.get('overview', ''), key=f"overview_{i}", placeholder="Summarize what your unit or command's mission was, and how your role contributed to it.")
            entry['responsibilities'] = st.text_area("Key Responsibilities", value=entry.get('responsibilities', ''), key=f"resp_{i}", placeholder="List your day-to-day responsibilities or duties. Paste bullets or freeform text.")
            entry['impact'] = st.text_area("Measurable Impact", value=entry.get('impact', ''), key=f"impact_{i}", placeholder="Include tangible accomplishments like promotions, saved time, improved performance, or cost savings.")

    with st.expander("Add Volunteer Experience"):
        org = st.text_input("Organization")
        role = st.text_input("Volunteer Role")
        dates = st.text_input("Dates (e.g. 2022 – Present)")
        impact = st.text_area("What did you do or achieve?")
        if st.button("➕ Add Volunteer Entry"):
            if org.strip() or role.strip() or dates.strip() or impact.strip():
                st.session_state.volunteer.append({"org": org.strip(), "role": role.strip(), "dates": dates.strip(), "impact": impact.strip()})
            else:
                st.warning("Please fill in at least one field before adding a volunteer entry.")

    for v in st.session_state.volunteer:
        if v['org'] or v['role'] or v['dates'] or v['impact']:
            st.markdown(f"**{v['role']}**, *{v['org']}* ({v['dates']})<br>{v['impact']}", unsafe_allow_html=True)

    st.session_state.skills = st.text_area("Skills (comma-separated)", value=st.session_state.get("skills", ""), placeholder="e.g., Cybersecurity, RMF, Azure, Python, Leadership")

    if st.button("← Back"):
        st.session_state.step = 1
    if st.button("Next →", key="go_to_step_3"):
        st.session_state.step = 3

# Step 3 - Review & Download (Updated to include Education and Certifications)
elif st.session_state.step == 3:
    st.subheader("Step 3 of 3: Review & Download")
    st.session_state.job_target = st.text_input("What job are you targeting?", value=st.session_state.get("job_target", ""))

    doc = Document()
    doc.add_heading(st.session_state.full_name, 0)
    doc.add_paragraph(f"Email: {st.session_state.email} | Phone: {st.session_state.phone} | LinkedIn: {st.session_state.linkedin}")
    doc.add_paragraph(f"Target Role: {st.session_state.job_target}")
    doc.add_paragraph("")

    doc.add_heading("Summary", level=1)
    summary = f"Military professional with {st.session_state.yos} years of experience as a {st.session_state.rank} ({st.session_state.mos}) with {st.session_state.clearance} clearance. Skilled in {st.session_state.job_target.lower()}, operational leadership, and mission-focused execution."
    doc.add_paragraph(summary)

    doc.add_heading("Professional Experience", level=1)
    sorted_jobs = sorted([e for e in st.session_state.job_data[:st.session_state.positions] if any(e.values())], key=extract_date, reverse=True)
    for entry in sorted_jobs:
        doc.add_heading(entry.get('title', ''), level=2)
        doc.add_paragraph(f"{entry.get('unit', '')} | {entry.get('location', '')} | {entry.get('dates', '')}")
        doc.add_paragraph(f"Mission: {clean_input(entry.get('overview', ''))}")
        doc.add_paragraph(f"Responsibilities: {clean_input(entry.get('responsibilities', ''))}")
        doc.add_paragraph(f"Impact: {clean_input(entry.get('impact', ''))}")

    # Add education
    remove_section(doc, "Education")
    if st.session_state.education:
        add_bold_paragraph("Education")
        doc.add_paragraph(st.session_state.education)

    # Add certifications
    remove_section(doc, "Certifications")
    if st.session_state.certifications:
        add_bold_paragraph("Certifications")
        for cert in st.session_state.certifications.strip().split("\n"):
            doc.add_paragraph(f"• {cert}")

    doc.add_heading("Awards", level=1)
    for award in st.session_state.awards.strip().split("\n"):
        doc.add_paragraph(f"• {award.strip()}")

    if st.session_state.volunteer:
        doc.add_heading("Volunteer Experience", level=1)
        for v in st.session_state.volunteer:
            if v['org'] or v['role'] or v['dates'] or v['impact']:
                doc.add_heading(v['role'], level=2)
                doc.add_paragraph(f"{v['org']} | {v['dates']}")
                doc.add_paragraph(v['impact'])

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join([s.strip() for s in st.session_state.skills.split(",") if s.strip()]))

    if st.button("← Back", key="back_to_step_2"):
        st.session_state.step = 2
    if st.button("🎯 Generate My Resume (.docx)"):
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("📄 Click here to download your resume", buffer, file_name="resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
