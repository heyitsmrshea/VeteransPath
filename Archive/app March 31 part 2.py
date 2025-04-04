import streamlit as st
from resume_helper import translate_fitrep_to_resume, generate_professional_summary, generate_position_summary
import base64
import docx
import json
from io import BytesIO
from datetime import datetime
import os

# --- Set page config and app title ---
st.set_page_config(page_title="OperationMOS", layout="wide")
st.title("Welcome to OperationMOS")
st.subheader("Your Personal Transition Assistant")

# --- Session state ---
def init_session():
    keys = [
        "form_submitted", "bullets", "follow_up_question", "original_input",
        "target_role", "selected_group", "feedback_given", "clarification_history",
        "branch", "mos", "rank", "years", "executive_summary", "position_summary"
    ]
    for key in keys:
        if key not in st.session_state:
            st.session_state[key] = [] if key == "clarification_history" else "" if key in ["selected_group", "follow_up_question", "original_input", "target_role", "branch", "mos", "rank", "executive_summary", "position_summary"] else False if "_given" in key else 0 if key == "years" else [] if "bullets" in key else None
init_session()

# --- Sidebar Progress Tracker ---
with st.sidebar:
    st.markdown("### 📌 Progress")
    st.markdown("✅ Step 1: Intake")
    if st.session_state.form_submitted:
        st.markdown("✅ Step 2: Bullet Generation")
        if st.session_state.bullets:
            st.markdown("✅ Step 3: Finalize + Download")

# --- Intake Form ---
if not st.session_state.form_submitted:
    with st.form("background_form"):
        st.markdown("### 🎖️ Tell us about your service")
        st.session_state.branch = st.selectbox("Branch of Service", ["Army", "Navy", "Marine Corps", "Air Force", "Coast Guard", "Space Force"])
        st.session_state.mos = st.text_input("MOS / Job Code")
        st.session_state.rank = st.text_input("Last Rank Held")
        st.session_state.years = st.number_input("Years of Service", min_value=0, step=1)
        goal = st.selectbox("What's your goal?", [
            "Build my resume",
            "Prepare for interviews",
            "Map to civilian career",
            "Explore networking tips"
        ])
        submitted = st.form_submit_button("Continue")

    if submitted:
        st.session_state.form_submitted = True
        st.session_state.goal = goal
        st.success(f"Let's help you {goal.lower()}!")

# --- Resume Builder ---
if st.session_state.form_submitted and st.session_state.goal == "Build my resume":
    st.header("📝 Resume Builder")
    col1, col2 = st.columns([2, 1])

    with col1:
        if not st.session_state.follow_up_question:
            st.markdown("### 🧠 Describe your experience")

            high_level = st.text_area("1. What was your role or mission context?")
            responsibilities = st.text_area("2. What were your primary duties or technologies used?")
            impact = st.text_area("3. What outcomes or impact did your work have? (include metrics if possible)")
            target_role = st.text_input("🎯 Target Role (e.g., Cybersecurity Analyst)")

            if st.button("Generate Resume Bullets"):
                if high_level and responsibilities and impact and target_role:
                    full_input = f"""
High-Level Overview:
{high_level}

Roles and Responsibilities:
{responsibilities}

Impact and Outcomes:
{impact}
"""
                    st.session_state.original_input = full_input
                    st.session_state.target_role = target_role

                    # Generate resume bullets
                    result = translate_fitrep_to_resume(full_input, target_role)
                    if "follow_up" in result:
                        st.session_state.follow_up_question = result["follow_up"]
                        st.session_state.clarification_history.append(result["follow_up"])
                    elif "bullets" in result:
                        st.session_state.bullets = result["bullets"]
                        st.session_state.selected_group = "Bullet Group 1"

                        # Generate GPT-based executive summary
                        st.session_state.executive_summary = generate_professional_summary(
                            branch=st.session_state.branch,
                            rank=st.session_state.rank,
                            mos=st.session_state.mos,
                            years=st.session_state.years,
                            target_role=st.session_state.target_role
                        )

                        # Generate position summary block
                        st.session_state.position_summary = generate_position_summary(
                            mos=st.session_state.mos,
                            branch=st.session_state.branch,
                            years=st.session_state.years,
                            target_role=st.session_state.target_role
                        )
                else:
                    st.warning("Please complete all fields.")
        else:
            st.warning(f"🔍 GPT needs more detail: {st.session_state.follow_up_question}")
            follow_up_answer = st.text_area("Your answer:")

            if st.button("Submit Clarification"):
                combined_input = (
                    f"{st.session_state.original_input}\n\n"
                    f"Clarification History:\n{chr(10).join(st.session_state.clarification_history)}\n\n"
                    f"Most Recent Answer:\n{follow_up_answer}"
                )
                result = translate_fitrep_to_resume(combined_input, st.session_state.target_role)
                if "bullets" in result:
                    st.session_state.bullets = result["bullets"]
                    st.session_state.follow_up_question = ""
                elif "follow_up" in result:
                    st.session_state.follow_up_question = result["follow_up"]
                    st.session_state.clarification_history.append(result["follow_up"])

    with col2:
        st.markdown("### 📄 Live Preview")
        st.markdown(f"**Overview:** {st.session_state.get('original_input', '').split('Roles and Responsibilities:')[0].strip()}")
        st.markdown(f"**Target Role:** {st.session_state.get('target_role', '')}")

    if st.session_state.bullets:
        st.subheader("✅ Your Structured Resume Content")
        st.session_state.selected_group = st.selectbox("Choose your favorite bullet group:", options=[f"Bullet Group {i+1}" for i in range(len(st.session_state.bullets))])

        index = int(st.session_state.selected_group.split()[-1]) - 1
        selected_bullets = st.session_state.bullets[index]

        if isinstance(selected_bullets, list) and len(selected_bullets) == 3:
            st.markdown(f"- {selected_bullets[0]}")
            st.markdown(f"- {selected_bullets[1]}")
            st.markdown(f"- {selected_bullets[2]}")

            if st.button("💾 Save this resume session"):
                session_output = {
                    "timestamp": str(datetime.now()),
                    "target_role": st.session_state.target_role,
                    "bullets": selected_bullets
                }
                os.makedirs("saved_resumes", exist_ok=True)
                with open("saved_resumes/final_resume.json", "w") as f:
                    json.dump(session_output, f)
                st.success("Resume session saved.")

            st.markdown("---")
            st.markdown("### 💬 Was this resume translation accurate?")
            colf1, colf2 = st.columns(2)
            if colf1.button("👍 Yes, it worked!"):
                st.session_state.feedback_given = True
                st.success("Awesome! Let’s move to final download.")
            if colf2.button("👎 No, needs revision"):
                st.session_state.feedback_given = True
                st.warning("Got it. Please continue refining with follow-up prompts above.")

            st.markdown("---")
            st.header("📥 Download Final Resume")

            if st.button("Download .docx Resume"):
                try:
                    doc = docx.Document()
                    doc.add_heading("Professional Summary", level=1)
                    doc.add_paragraph(st.session_state.executive_summary)
                    doc.add_paragraph("")

                    doc.add_paragraph(st.session_state.position_summary)
                    doc.add_paragraph("")

                    doc.add_paragraph(selected_bullets[0])
                    doc.add_paragraph(selected_bullets[1])
                    doc.add_paragraph(selected_bullets[2])

                    buf = BytesIO()
                    doc.save(buf)
                    st.download_button("📄 Download .docx", data=buf.getvalue(), file_name="OperationMOS_Resume.docx")

                except Exception as e:
                    st.error(f"Something went wrong: {e}")
        else:
            st.error("❌ Invalid bullet group format. Expected a list of 3 items.")
